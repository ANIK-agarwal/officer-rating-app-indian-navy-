#-------------------------------------------------------------------------------
# Name:        module1
# Purpose:    creation V334321
#
# Author:      anik
#
# Created:     08-06-2025
# Copyright:   (c) anika 2025
# Licence:     <your licence> private
#-------------------------------------------------------------------------------

import os
import pandas as pd
from tkinter import *
from tkinter import ttk, messagebox
from docx import Document
from PIL import Image, ImageTk

# ---------- Load qualities from Excel ----------
def load_qualities_from_excel():
    df = pd.read_excel("qualities.xlsx")
    df.columns = df.columns.str.strip()  # remove extra spaces
    quality_dict = {}

    print("Loaded keys:", quality_dict.keys())
    print("Excel Columns:", df.columns.tolist())  # Debug print

    for i, row in df.iterrows():
        try:
            topic = str(row['Topic']).strip()
            sub_topic = str(row['Sub-Topic']).strip()
            rating = str(row['Rating']).strip()
            quality = str(row['Quality Description']).strip()

            key = f"{topic}__{sub_topic}__{rating}"
            quality_dict[key] = quality
        except Exception as e:
            print(f"Error processing row {i}: {e}")

    return quality_dict


# ---------- Main GUI ----------
def build_app():
    quality_dict = load_qualities_from_excel()

    root = Tk()
    root.title("Officer Quality Rating - SHADE CARD")
    root.geometry("1100x750")

    # Create a Canvas inside the root
    main_container = Frame(root)
    main_container.pack(fill=BOTH, expand=True)

    #creating scroolbar inside canvas
    v_scrollbar = Scrollbar(main_container, orient=VERTICAL)
    v_scrollbar.pack(side=RIGHT, fill=Y)

    main_canvas = Canvas(main_container, yscrollcommand=v_scrollbar.set, bg="#001f3f")
    main_canvas.pack(side=LEFT, fill=BOTH, expand=True)

    v_scrollbar.config(command=main_canvas.yview)


    # Create a scrollable Frame inside the Canvas
    scroll_frame = Frame(main_canvas, bg="#001f3f")
    scroll_frame.bind(
        "<Configure>",
        lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all"))
    )
    main_canvas.create_window((0, 0), window=scroll_frame, anchor="nw")
    scroll_frame.bind("<Configure>", lambda e: main_canvas.configure(scrollregion=main_canvas.bbox("all")))


    # -------------------------- UI CONTENT BELOW --------------------------
    # ---------- Top Header ----------
    top_frame = Frame(scroll_frame, bg="#001f3f")
    top_frame.pack(fill=X, pady=10)

    title = Label(top_frame, text="SHADE CARD (officer quality rating)",
                  font=("Arial", 20, "bold"), fg="gold", bg="#001f3f")
    title.pack(side=LEFT, padx=20)

    try:
        img = Image.open("navy_logo.png")
        img = img.resize((100, 100))
        logo_img = ImageTk.PhotoImage(img)
        Label(top_frame, image=logo_img, bg="#001f3f").pack(side=RIGHT, padx=20)
    except:
        Label(top_frame, text="[Logo Missing]", fg="white", bg="#001f3f").pack(side=RIGHT, padx=20)

    # ---------- User Info ----------
    user_frame = Frame(scroll_frame, bg="#001f3f")
    user_frame.pack(pady=5)
    Label(user_frame, text="Name:", fg="white", bg="#001f3f").pack(side=LEFT)
    name_var = StringVar()
    Entry(user_frame, textvariable=name_var, width=30).pack(side=LEFT, padx=10)

    # radio buttons selection
    Label(user_frame, text="Gender:", fg="white", bg="#001f3f").pack(side=LEFT, padx=(10, 5))
    gender_var = StringVar(value="Male")  # default value

    Radiobutton(user_frame, text="Male", variable=gender_var, value="Male", bg="#001f3f", fg="white", selectcolor="#001f3f").pack(side=LEFT)
    Radiobutton(user_frame, text="Female", variable=gender_var, value="Female", bg="#001f3f", fg="white", selectcolor="#001f3f").pack(side=LEFT)

    # ---------- Topics ----------
    topics = [
        ('Planning & Organizing', ['Effective Intelligence', 'Reasoning Ability', 'Organising Ability','Power of Expression']),
        ('Social Effectiveness', ['Initiative', 'Self Confidence', 'Speed of Decision', 'Ability to Infiuence people', 'Liveliness']),
        ('Social adjustments', ['Co-operatiom', 'Social Adeptability', 'Sense of Responsibility']),
        ('Dynamic', ['Determination', 'Stamina', 'Courage'])
    ]

    input_vars = {}
    preview_text = StringVar()

    def on_rating_change(main_topic, sub_topic, v1, v2, v3):
        key = f"{main_topic}__{sub_topic}__{v1.get().strip().upper()}{v2.get().strip().upper()}{v3.get().strip().upper()}"
        quality = quality_dict.get(key, "(No quality mapped)")
        preview_text.set(f"{quality}")

    form_frame = Frame(scroll_frame, bg="#001f3f")
    form_frame.pack(padx=10, pady=10, fill=BOTH, expand=True)

    for i in range(2):
        row_frame = Frame(form_frame, bg="#001f3f")
        row_frame.pack(fill=X, pady=5)

        for j in range(2):
            index = i * 2 + j
            if index >= len(topics): continue
            main_topic, sub_topics = topics[index]

            topic_frame = Frame(row_frame, bg="#003366", bd=2, relief=RIDGE)
            topic_frame.pack(side=LEFT, padx=10, pady=10, ipadx=10, ipady=10, fill=BOTH, expand=True)

            Label(topic_frame, text=main_topic, font=("Arial", 12, "bold"), fg="white", bg="#003366").pack()

            input_vars[main_topic] = []
            for sub in sub_topics:
                Label(topic_frame, text=sub, fg="white", bg="#003366").pack(anchor="w")
                f = Frame(topic_frame, bg="#003366")
                f.pack(anchor="w", pady=2)
                v1, v2, v3 = StringVar(), StringVar(), StringVar()
                combo1 = ttk.Combobox(f, textvariable=v1, width=3, values=[str(x) for x in range(1, 11)])
                combo2 = ttk.Combobox(f, textvariable=v2, width=3, values=list("ABCDEFGHIJ"))
                combo3 = ttk.Combobox(f, textvariable=v3, width=3, values=list("ABCDEFGHIJ"))
                combo1.pack(side=LEFT)
                combo2.pack(side=LEFT, padx=2)
                combo3.pack(side=LEFT)

                def bind_change(v1=v1, v2=v2, v3=v3, m=main_topic, s=sub):
                    v1.trace_add("write", lambda *_: on_rating_change(m, s, v1, v2, v3))
                    v2.trace_add("write", lambda *_: on_rating_change(m, s, v1, v2, v3))
                    v3.trace_add("write", lambda *_: on_rating_change(m, s, v1, v2, v3))

                bind_change()
                input_vars[main_topic].append((sub, v1, v2, v3))


   # ---------- Preview (with border and custom background) ----------
    preview_frame = Frame(scroll_frame, bg="#001f3f")
    preview_frame.pack(pady=10, fill=X)

    bordered_preview = Frame(preview_frame, bg="#004080", bd=3, relief=GROOVE)
    bordered_preview.pack(padx=20, pady=5, fill=X, expand=True)

   # adding min width
    bordered_preview.pack_propagate(False)
    bordered_preview.config(width=900, height=120)

    Label(bordered_preview, text="Live Preview", font=("Arial", 12, "bold"),
     fg="white", bg="#004080").pack(anchor="w", padx=10, pady=(5, 0))

    Label(bordered_preview, textvariable=preview_text, wraplength=600,
          justify=LEFT, fg="white", bg="#004080", font=("Arial", 11)).pack(padx=10, pady=5, anchor="w")

    # ---------- Buttons ----------
    def clear_all():
        name_var.set("")
        gender_var.set("")
        preview_text.set("")
        for main in input_vars:
            for _, v1, v2, v3 in input_vars[main]:
                v1.set("")
                v2.set("")
                v3.set("")

    def generate_doc():
        doc = Document()
        doc.add_heading(f"SHADE CARD REPORT", 0)
        doc.add_paragraph(f"Name: {name_var.get()}    Gender: {gender_var.get()}")

        for main_topic in input_vars:
            doc.add_heading(main_topic, level=1)
            text_block = ""
            for sub, v1, v2, v3 in input_vars[main_topic]:
                key = f"{main_topic}__{sub}__{v1.get().strip().upper()}{v2.get().strip().upper()}{v3.get().strip().upper()}"
                quality = quality_dict.get(key, "(No quality mapped)")
                text_block += quality + ". "
            doc.add_paragraph(text_block.strip())

        doc.save("Officer_Rating_Report.docx")
        messagebox.showinfo("Done", "Word document has been saved.")

    def confirm_and_generate():
        review = ""
        for main in input_vars:
            review += f"\n{main}:\n"
            for sub, v1, v2, v3 in input_vars[main]:
                review += f"  {sub}: {v1.get()} {v2.get()} {v3.get()}\n"

        result = messagebox.askquestion("Confirm Submission",
                    f"Please review the ratings: \n{review}\n\nPress 'Yes' to Continue or 'No' to Modify")
        if result == 'yes':
            generate_doc()

    button_frame = Frame(scroll_frame, bg="#001f3f")
    button_frame.pack(pady=10)

    # continue formation
    Button(button_frame, text="Submit", command=confirm_and_generate).pack(side=LEFT, padx=10)
    Button(button_frame, text="Clear", command=clear_all).pack(side=LEFT, padx=10)
    Button(button_frame, text="Close", command=root.destroy).pack(side=LEFT, padx=10)

    # ---------- Instructions (Moved Below Buttons) ----------
    instructions = [
        "1. Fill name and select gender.",
        "2. Select ratings for each sub-topic.",
        "3. Preview updates live below.",
        "4. Click Submit to review all entries.",
        "5. Modify or Continue to generate report."
    ]

    instr_frame = Frame(scroll_frame, bg="#001f3f")
    instr_frame.pack(pady=10)

    Label(instr_frame, text="Instructions", font=("Arial", 12, "bold"), fg="gold", bg="#001f3f").pack(anchor="w", pady=(5, 2))
    for line in instructions:
        Label(instr_frame, text=line, fg="white", bg="#001f3f").pack(anchor="w")


    # using mouse for scroll
    main_canvas.bind_all("<MouseWheel>", lambda e: main_canvas.yview_scroll(int(-1*(e.delta/120)), "units"))

    # main loop
    root.mainloop()

# Run the app
build_app()

